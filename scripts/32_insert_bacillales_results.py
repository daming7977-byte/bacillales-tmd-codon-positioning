from docx import Document
from copy import deepcopy

DOCX = "/Users/liming/Desktop/秋山研/2026李明投稿论文/李明博论/MBE_submission_manuscript_v1.2.docx"

doc = Document(DOCX)

heading_text = "Independent analysis in Bacillales reveals partial and anchor-specific support"

body = [
    (
        "To assess whether the TMD-relative organization of low-adaptation synonymous codon "
        "segments extends beyond Enterobacterales, I performed an independent analysis across "
        "10 representative Bacillales genomes. Strict one-to-one orthologous groups were subjected "
        "to a conservative topology-quality-control workflow, yielding 153 topology-qualified "
        "multi-pass membrane-protein families. For the primary positional analysis, proteins were "
        "further restricted to those whose predicted TMD number matched the modal TMD count of the "
        "corresponding orthogroup, retaining 1,372 proteins. Among 1,086 topology-qualified homologous "
        "TMD clusters, 994 remained represented in at least eight species after this filter."
    ),
    (
        "Using the same species-specific low-adaptation definition, a minimum segment length of "
        "three consecutive codons, nearest-TMD assignment framework, and 1,000 within-protein "
        "synonymous-codon permutations, 30 TMD-start units and 33 TMD-end units satisfied the "
        "requirement of an assigned low-adaptation segment in at least three species. "
        "TMD-start-relative positional variance showed no reduction relative to the synonymous null: "
        "the observed median variance was 2,031.10 aa² compared with a null median of 1,817.13 aa² "
        "(empirical one-sided P = 0.595). In contrast, TMD-end-relative variance was directionally "
        "lower in the observed data, with a median of 739.08 aa² compared with a null median of "
        "1,693.00 aa², although this reduction did not reach permutation-based statistical significance "
        "(P = 0.0829). The supplementary TMD-center analysis likewise showed no evidence of reduced "
        "positional variance (1,769.78 versus a null median of 1,813.82 aa²; P = 0.472)."
    ),
    (
        "Thus, the strong TMD-boundary-relative constraint detected in Enterobacterales was not fully "
        "reproduced in Bacillales. Instead, the independent lineage showed partial, anchor-specific "
        "support concentrated at TMD ends, suggesting that the evolutionary organization of local "
        "synonymous-codon features relative to membrane topology may differ among bacterial lineages."
    ),
]

# Insert before Discussion
target = None
for p in doc.paragraphs:
    if p.text.strip() == "Discussion":
        target = p
        break

if target is None:
    raise RuntimeError('Could not find heading "Discussion"')

# create temporary paragraphs
h = doc.add_paragraph()
try:
    h.style = "Heading 2"
except KeyError:
    pass
h.add_run(heading_text)

new_paras = []
for text in body:
    p = doc.add_paragraph(text)
    new_paras.append(p)

# move them before Discussion
for p in [h] + new_paras:
    target._element.addprevious(deepcopy(p._element))

# remove temporary copies at end
for p in [h] + new_paras:
    el = p._element
    el.getparent().remove(el)

doc.save(DOCX)

print("Updated:", DOCX)
