from docx import Document
from copy import deepcopy

DOCX = "/Users/liming/Desktop/秋山研/2026李明投稿论文/李明博论/MBE_submission_manuscript_v1.2.docx"

HEADING = "Independent Bacillales replication analysis"
TARGET = "Sensitivity analysis excluding first TMDs"

doc = Document(DOCX)

# --------------------------------------------------
# Prevent accidental duplicate insertion
# --------------------------------------------------

for p in doc.paragraphs:
    if p.text.strip() == HEADING:
        print("Section already exists; no changes made.")
        raise SystemExit


# --------------------------------------------------
# Find insertion point
# --------------------------------------------------

target = None

for p in doc.paragraphs:
    if p.text.strip() == TARGET:
        target = p
        break

if target is None:
    raise RuntimeError(
        f'Could not find insertion target: "{TARGET}"'
    )


# --------------------------------------------------
# Bacillales Methods text
# --------------------------------------------------

body = [

    (
        "To evaluate the generalizability of the Enterobacterales result in a "
        "phylogenetically distinct bacterial lineage, I performed an independent "
        "analysis across 10 representative Bacillales genomes. Orthology was "
        "defined using reciprocal best hits under conservative sequence-similarity "
        "and coverage criteria, followed by a strict one-to-one family filter. "
        "Orthologous groups represented in at least eight of the 10 species were "
        "retained for topology analysis."
    ),

    (
        "Transmembrane topology was predicted using DeepTMHMM for all proteins in "
        "the retained orthologous groups. Candidate multi-pass membrane-protein "
        "families were defined as orthogroups with a modal TMD number of at least "
        "three. Protein sequences within each candidate family were aligned using "
        "MAFFT. Predicted TMD intervals were projected onto the corresponding "
        "protein alignments, and homologous TMDs were identified by alignment-"
        "coordinate overlap. TMDs from different proteins were connected when their "
        "aligned intervals overlapped by at least 50% of the shorter interval, and "
        "connected components were treated as candidate homologous TMD clusters."
    ),

    (
        "Homologous TMD clusters were retained when they were represented in at "
        "least 80% of family members, no protein contributed more than one TMD to "
        "the same cluster, and the median absolute deviations of both aligned "
        "TMD-start and TMD-end coordinates were no greater than two residues. "
        "Families were considered topology-qualified when they contained at least "
        "three such homologous TMD clusters and showed no violations of conserved "
        "N-to-C-terminal TMD order. This procedure yielded 153 topology-qualified "
        "Bacillales membrane-protein families comprising 1,432 proteins and 1,086 "
        "homologous TMD clusters."
    ),

    (
        "Species-specific synonymous codon weights were calculated independently "
        "for each Bacillales genome using the same relative-weight definition as "
        "in the Enterobacterales analysis. Low-adaptation thresholds were defined "
        "independently within each species as the frozen bottom decile of "
        "codon-position weights across the topology-qualified membrane-protein "
        "analysis set, with ties retained. Low-adaptation segments were called "
        "using the same minimum length of three consecutive codons, with initiation "
        "codons and stop codons excluded from classification."
    ),

    (
        "For the primary positional analysis, proteins were restricted to those "
        "whose predicted TMD number matched the modal TMD count of the corresponding "
        "orthogroup, as in the Enterobacterales analysis. This retained 1,372 "
        "proteins. Homologous TMD clusters were required to remain represented in "
        "at least eight species after this filter, yielding 994 eligible TMD "
        "clusters. Low-adaptation segments were assigned independently for "
        "TMD-start, TMD-end, and TMD-center analyses to the nearest corresponding "
        "TMD anchor within each protein. Assignment was performed before "
        "homologous-cluster eligibility filtering; if the nearest TMD belonged to "
        "a cluster that did not satisfy the frozen topology criteria, the segment "
        "was not reassigned to a more distant TMD. When multiple segments from the "
        "same protein mapped to the same TMD, only the segment with the smallest "
        "absolute anchor-relative distance was retained."
    ),

    (
        "Cross-species positional variance and unit eligibility were calculated "
        "using the same criteria as in the Enterobacterales analysis. Homologous "
        "TMD units were evaluated when represented in at least eight species, and "
        "a unit contributed to the positional-conservation statistic when at least "
        "three species contained an assigned low-adaptation segment. Sample variance "
        "of anchor-relative positions was calculated for each qualifying unit, and "
        "the median variance across units was used as the summary statistic."
    ),

    (
        "The synonymous-codon null was generated using 1,000 within-protein "
        "permutations. Within each coding sequence, synonymous codons were shuffled "
        "only among positions encoding the same amino acid, preserving the encoded "
        "protein sequence and amino-acid-specific codon composition. After each "
        "permutation, the frozen species-specific codon weights and thresholds were "
        "reapplied, low-adaptation segments were recalled de novo, and anchor-specific "
        "nearest-TMD assignment and unit qualification were repeated. Empirical "
        "one-sided P values were calculated as P = (k + 1)/(N + 1), where k was "
        "the number of permuted datasets with a median variance less than or equal "
        "to the observed statistic and N = 1,000. Random seed 20260825 was used."
    ),
]


# --------------------------------------------------
# Create section temporarily at end
# --------------------------------------------------

h = doc.add_paragraph()

# Match the style of the next Methods subsection
try:
    h.style = target.style
except Exception:
    pass

h.add_run(HEADING)

new_paragraphs = []

for text in body:
    p = doc.add_paragraph(text)
    new_paragraphs.append(p)


# --------------------------------------------------
# Move before sensitivity-analysis subsection
# --------------------------------------------------

for p in [h] + new_paragraphs:
    target._element.addprevious(
        deepcopy(p._element)
    )


# Remove temporary end-of-document copies
for p in [h] + new_paragraphs:
    el = p._element
    el.getparent().remove(el)


doc.save(DOCX)

print("Inserted Methods section:")
print(HEADING)
print()
print("Updated:")
print(DOCX)
