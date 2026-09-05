from docx import Document
from copy import deepcopy

DOCX = "/Users/liming/Desktop/秋山研/2026李明投稿论文/李明博论/MBE_submission_manuscript_v1.2.docx"

NEW_PARAGRAPH = (
    "The independent Bacillales analysis further defined the phylogenetic scope "
    "of this pattern. In contrast to the strong reduction in both TMD-start- and "
    "TMD-end-relative positional variance observed in Enterobacterales, the "
    "Bacillales analysis did not show reduced variance relative to TMD starts. "
    "TMD-end-relative variance was substantially lower than the median of the "
    "synonymous-permutation null distribution (739.08 versus 1,693.00 aa²), but "
    "this directional reduction did not reach permutation-based statistical "
    "significance (P = 0.0829). The TMD-center analysis likewise provided no "
    "evidence of reduced positional variance. Thus, the Enterobacterales pattern "
    "does not appear to represent a simple universally conserved rule across "
    "bacterial orders. Instead, the strength and anchor dependence of TMD-relative "
    "synonymous-codon organization may vary among evolutionary lineages. Because "
    "the Bacillales analysis contained only 30–33 qualifying TMD units depending "
    "on anchor type, the nonsignificant TMD-end trend should be interpreted "
    "cautiously and does not by itself establish a conserved mechanism."
)

doc = Document(DOCX)

# prevent duplication
for p in doc.paragraphs:
    if p.text.strip().startswith(
        "The independent Bacillales analysis further defined"
    ):
        print("Discussion paragraph already exists; no changes made.")
        raise SystemExit

# locate Discussion heading
discussion_index = None

for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "Discussion":
        discussion_index = i
        break

if discussion_index is None:
    raise RuntimeError('Could not find "Discussion" heading.')

# find first non-empty paragraph after Discussion
first_discussion_para = None

for p in doc.paragraphs[discussion_index + 1:]:
    if p.text.strip():
        first_discussion_para = p
        break

if first_discussion_para is None:
    raise RuntimeError("Could not find first Discussion paragraph.")

# create new paragraph at end temporarily
new_p = doc.add_paragraph(NEW_PARAGRAPH)

# copy body style from existing Discussion paragraph
try:
    new_p.style = first_discussion_para.style
except Exception:
    pass

# insert AFTER first Discussion paragraph
first_discussion_para._element.addnext(
    deepcopy(new_p._element)
)

# remove temporary copy
el = new_p._element
el.getparent().remove(el)

doc.save(DOCX)

print("Inserted Bacillales Discussion paragraph.")
print("Updated:", DOCX)
