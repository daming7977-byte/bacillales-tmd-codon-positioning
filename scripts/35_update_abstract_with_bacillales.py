from docx import Document

DOCX = "/Users/liming/Desktop/秋山研/2026李明投稿论文/李明博论/MBE_submission_manuscript_v1.2.docx"

OLD = (
    "A set of 197 matched soluble-protein families provided additional but "
    "anchor-dependent evidence, with stronger support for TMD-start-relative "
    "than TMD-end-relative constraint."
)

NEW = (
    "A set of 197 matched soluble-protein families provided additional but "
    "anchor-dependent evidence, with stronger support for TMD-start-relative "
    "than TMD-end-relative constraint. An independent analysis across 10 "
    "Bacillales genomes did not reproduce the TMD-start signal, but showed a "
    "directional reduction in TMD-end-relative variance that did not reach "
    "permutation-based statistical significance (P = 0.0829)."
)

doc = Document(DOCX)

# prevent accidental duplicate insertion
for p in doc.paragraphs:
    if (
        "An independent analysis across 10 Bacillales genomes"
        in p.text
    ):
        print("Bacillales Abstract sentence already exists; no changes made.")
        raise SystemExit

found = False

for p in doc.paragraphs:
    if OLD in p.text:
        p.text = p.text.replace(OLD, NEW)
        found = True
        break

if not found:
    raise RuntimeError(
        "Could not find the expected Abstract sentence."
    )

doc.save(DOCX)

print("Updated Abstract with Bacillales result.")
print("Saved:", DOCX)
