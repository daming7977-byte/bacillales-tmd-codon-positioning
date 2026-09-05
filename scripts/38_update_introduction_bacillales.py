from docx import Document

DOCX = "/Users/liming/Desktop/秋山研/2026李明投稿论文/李明博论/MBE_submission_manuscript_v1.2.docx"

OLD = (
    "Matched soluble-family analyses provided additional but anchor-dependent "
    "evidence, with stronger support for TMD-start-relative constraint than for "
    "TMD-end-relative constraint. These results reveal evolutionarily constrained "
    "positioning of low-adaptation synonymous codon segments relative to homologous "
    "TMD boundaries, while leaving open whether this pattern directly reflects "
    "modulation of translation kinetics, membrane insertion, cotranslational folding, "
    "or another selective constraint."
)

NEW = (
    "Matched soluble-family analyses provided additional but anchor-dependent "
    "evidence, with stronger support for TMD-start-relative constraint than for "
    "TMD-end-relative constraint. I additionally tested the generalizability of "
    "this pattern in an independent set of 10 Bacillales genomes using the same "
    "codon-segment definition, positional statistic, and synonymous-permutation "
    "framework. These results reveal evolutionarily constrained positioning of "
    "low-adaptation synonymous codon segments relative to homologous TMD boundaries "
    "in Enterobacterales, while also defining the phylogenetic limits of this pattern "
    "and leaving open whether it directly reflects modulation of translation kinetics, "
    "membrane insertion, cotranslational folding, or another selective constraint."
)

doc = Document(DOCX)

# avoid duplicate insertion
for p in doc.paragraphs:
    if "I additionally tested the generalizability of this pattern" in p.text:
        print("Introduction already updated; no changes made.")
        raise SystemExit

found = False

for p in doc.paragraphs:
    if OLD in p.text:
        p.text = p.text.replace(OLD, NEW)
        found = True
        break

if not found:
    raise RuntimeError("Could not find expected Introduction paragraph.")

doc.save(DOCX)

print("Introduction updated.")
print("Saved:", DOCX)
