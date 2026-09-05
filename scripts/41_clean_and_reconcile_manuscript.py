from docx import Document

DOCX = "/Users/liming/Desktop/秋山研/2026李明投稿论文/李明博论/MBE_submission_manuscript_v1.2.docx"

doc = Document(DOCX)

# --------------------------------------------------
# 1. mechanical spacing fixes detected by audit
# --------------------------------------------------

spacing_fixes = {
    "unclear.Here": "unclear. Here",
    "TMDs,under": "TMDs, under",
    "amongbacterial": "among bacterial",
    "153topology-qualified": "153 topology-qualified",
    "werefurther": "were further",
    "TMD-startunits": "TMD-start units",
    "medianof": "median of",
    "establisha": "establish a",
    "using thesame": "using the same",
    "1,372proteins": "1,372 proteins",
    "thatdid": "that did",
    "species, anda": "species, and a",
    "codonorganization": "codon organization",
    "Proteinorthology": "Protein orthology",
    "comparative andobservational": "comparative and observational",
    "testablefeature": "testable feature",
    "wereused": "were used",
    "Cross-speciespositional": "Cross-species positional",
    "pseudo-anchorcontrols": "pseudo-anchor controls",
    "availableat": "available at",
    "anobserved": "an observed",
    "replicate,one": "replicate, one",
}

spacing_changes = 0

for p in doc.paragraphs:
    old = p.text
    new = old

    for a, b in spacing_fixes.items():
        new = new.replace(a, b)

    if new != old:
        p.text = new
        spacing_changes += 1


# --------------------------------------------------
# 2. rename Bacillales Methods heading
# --------------------------------------------------

for p in doc.paragraphs:
    if p.text.strip() == "Independent Bacillales replication analysis":
        p.text = "Independent Bacillales analysis"


# --------------------------------------------------
# 3. soften Bacillales Results heading
# --------------------------------------------------

for p in doc.paragraphs:
    if p.text.strip() == (
        "Independent analysis in Bacillales reveals partial and anchor-specific support"
    ):
        p.text = (
            "Independent analysis in Bacillales reveals an anchor-specific directional pattern"
        )


# --------------------------------------------------
# 4. soften concluding Bacillales Results wording
# --------------------------------------------------

OLD_RESULTS = (
    "Thus, the strong TMD-boundary-relative constraint detected in "
    "Enterobacterales was not fully reproduced in Bacillales. Instead, the "
    "independent lineage showed partial, anchor-specific support concentrated "
    "at TMD ends, suggesting that the evolutionary organization of local "
    "synonymous-codon features relative to membrane topology may differ among "
    "bacterial lineages."
)

NEW_RESULTS = (
    "Thus, the strong TMD-boundary-relative constraint detected in "
    "Enterobacterales was not fully reproduced in Bacillales. Instead, the "
    "independent lineage showed an anchor-specific directional pattern at TMD "
    "ends, suggesting that the strength and anchor dependence of local "
    "synonymous-codon organization relative to membrane topology may differ "
    "among bacterial lineages."
)

for p in doc.paragraphs:
    if OLD_RESULTS in p.text:
        p.text = p.text.replace(OLD_RESULTS, NEW_RESULTS)


# --------------------------------------------------
# 5. update outdated Discussion limitation
# --------------------------------------------------

OLD_LIMITATION = (
    "Second, the analysis was restricted to Enterobacterales and to strict "
    "one-to-one orthologous families with sufficiently conserved multi-pass "
    "membrane topology. The extent to which the same positional constraint "
    "applies across more deeply diverged bacterial lineages remains unknown."
)

NEW_LIMITATION = (
    "Second, the primary analysis was centered on Enterobacterales, whereas "
    "the independent Bacillales analysis provided only a single additional "
    "order-level comparison and contained fewer qualifying TMD units. The "
    "extent to which TMD-relative synonymous-codon positional constraint is "
    "shared across broader bacterial diversity therefore remains unresolved."
)

for p in doc.paragraphs:
    if OLD_LIMITATION in p.text:
        p.text = p.text.replace(OLD_LIMITATION, NEW_LIMITATION)


doc.save(DOCX)

print("Saved:", DOCX)
print("Paragraphs with spacing corrections:", spacing_changes)
