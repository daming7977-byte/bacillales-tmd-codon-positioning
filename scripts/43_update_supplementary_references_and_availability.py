from docx import Document

DOCX = "/Users/liming/Desktop/秋山研/2026李明投稿论文/李明博论/MBE_submission_manuscript_v1.2.docx"

doc = Document(DOCX)

changes = 0

# --------------------------------------------------
# 1. Bacillales Results: add supplementary references
# --------------------------------------------------

repls = {

    "yielding 153 topology-qualified multi-pass membrane-protein families. "
    "For the primary positional analysis, proteins were further restricted to "
    "those whose predicted TMD number matched the modal TMD count of the "
    "corresponding orthogroup, retaining 1,372 proteins. Among 1,086 "
    "topology-qualified homologous TMD clusters, 994 remained represented in "
    "at least eight species after this filter.":

    "yielding 153 topology-qualified multi-pass membrane-protein families "
    "(Supplementary Table B2). For the primary positional analysis, proteins "
    "were further restricted to those whose predicted TMD number matched the "
    "modal TMD count of the corresponding orthogroup, retaining 1,372 proteins "
    "(Supplementary Table B4). Among 1,086 topology-qualified homologous TMD "
    "clusters (Supplementary Table B3), 994 remained represented in at least "
    "eight species after this filter.",


    "(1,769.78 versus a null median of 1,813.82 aa²; P = 0.472).":

    "(1,769.78 versus a null median of 1,813.82 aa²; P = 0.472). "
    "Full unit-level and permutation-null results are provided in "
    "Supplementary Tables B7–B9.",
}

for p in doc.paragraphs:
    old = p.text
    new = old

    for a, b in repls.items():
        new = new.replace(a, b)

    if new != old:
        p.text = new
        changes += 1


# --------------------------------------------------
# 2. Bacillales Methods: species table reference
# --------------------------------------------------

for p in doc.paragraphs:

    if (
        p.text.startswith(
            "To evaluate the generalizability of the Enterobacterales result"
        )
        and "Supplementary Table B1" not in p.text
    ):
        p.text = p.text.replace(
            "Orthologous groups represented in at least eight of the 10 species "
            "were retained for topology analysis.",
            "Orthologous groups represented in at least eight of the 10 species "
            "were retained for topology analysis. The Bacillales species panel "
            "and assembly accessions are listed in Supplementary Table B1."
        )
        changes += 1
        break


# --------------------------------------------------
# 3. Bacillales Methods: threshold + segment references
# --------------------------------------------------

for p in doc.paragraphs:

    if (
        p.text.startswith(
            "Species-specific synonymous codon weights were calculated "
            "independently for each Bacillales genome"
        )
    ):

        if "Supplementary Table B5" not in p.text:
            p.text = p.text.replace(
                "with ties retained.",
                "with ties retained (Supplementary Table B5)."
            )

        if "Supplementary Table B6" not in p.text:
            p.text = p.text.replace(
                "with initiation codons and stop codons excluded from "
                "classification.",
                "with initiation codons and stop codons excluded from "
                "classification. The resulting low-adaptation segments are "
                "provided in Supplementary Table B6."
            )

        changes += 1
        break


# --------------------------------------------------
# 4. Bacillales Methods: primary unit references
# --------------------------------------------------

for p in doc.paragraphs:

    if p.text.startswith(
        "For the primary positional analysis, proteins were restricted"
    ):

        if "Supplementary Table B4" not in p.text:
            p.text = p.text.replace(
                "This retained 1,372 proteins.",
                "This retained 1,372 proteins "
                "(Supplementary Table B4)."
            )

        if "Supplementary Table B3" not in p.text:
            p.text = p.text.replace(
                "yielding 994 eligible TMD clusters.",
                "yielding 994 eligible TMD clusters "
                "(Supplementary Table B3)."
            )

        changes += 1
        break


for p in doc.paragraphs:

    if p.text.startswith(
        "Cross-species positional variance and unit eligibility were calculated"
    ):

        if "Supplementary Table B7" not in p.text:
            p.text += (
                " Unit-level primary results are provided in "
                "Supplementary Table B7."
            )

        changes += 1
        break


for p in doc.paragraphs:

    if p.text.startswith(
        "The synonymous-codon null was generated using 1,000 within-protein"
    ):

        if "Supplementary Tables B8 and B9" not in p.text:
            p.text += (
                " The full permutation-null summary and final observed-versus-null "
                "statistics are provided in Supplementary Tables B8 and B9."
            )

        changes += 1
        break


# --------------------------------------------------
# 5. Update Data Availability
# --------------------------------------------------

data_anchor = (
    "Processed analysis tables underlying the main figures and statistical analyses"
)

for p in doc.paragraphs:

    if p.text.startswith(data_anchor):

        if "Bacillales" not in p.text:

            p.text += (
                " Processed tables underlying the independent Bacillales analysis "
                "are provided as Supplementary Tables B1–B9."
            )

        changes += 1
        break


# --------------------------------------------------
# 6. Update Code Availability
# --------------------------------------------------

code_anchor = (
    "Custom scripts used for ortholog filtering, topology integration"
)

for p in doc.paragraphs:

    if p.text.startswith(code_anchor):

        if "Bacillales" not in p.text:

            p.text += (
                " Custom scripts for the independent Bacillales orthology, "
                "topology clustering, codon-weight calculation, low-adaptation "
                "segment detection, TMD-relative assignment, and synonymous-"
                "permutation analyses will be deposited with the final "
                "reproducibility release."
            )

        changes += 1
        break


doc.save(DOCX)

print("Saved:", DOCX)
print("Paragraphs updated:", changes)
