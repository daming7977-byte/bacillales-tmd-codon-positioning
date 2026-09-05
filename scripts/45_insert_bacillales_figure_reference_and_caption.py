from docx import Document

DOCX = "/Users/liming/Desktop/秋山研/2026李明投稿论文/李明博论/MBE_submission_manuscript_v1.2.docx"

doc = Document(DOCX)

changes = 0

# --------------------------------------------------
# 1. Update Bacillales Results reference
# --------------------------------------------------

OLD = (
    "Full unit-level and permutation-null results are provided in "
    "Supplementary Tables B7–B9."
)

NEW = (
    "Full unit-level and permutation-null results are provided in "
    "Supplementary Figure B1 and Supplementary Tables B7–B9."
)

for p in doc.paragraphs:
    if OLD in p.text:
        p.text = p.text.replace(OLD, NEW)
        changes += 1
        break


# --------------------------------------------------
# 2. Add Supplementary Figure B1 legend
#    Insert before Data Availability if possible
# --------------------------------------------------

CAPTION_HEAD = "Supplementary Figure B1."

CAPTION = (
    "Supplementary Figure B1. Synonymous-permutation null distributions for "
    "the independent Bacillales analysis. Distributions of the median "
    "cross-species variance in low-adaptation-segment position relative to "
    "homologous TMD starts, ends, and centers across 1,000 within-protein "
    "synonymous-codon permutations. Solid vertical lines indicate the observed "
    "median variance and dashed vertical lines indicate the median of the "
    "corresponding permutation-null distribution. The Bacillales analysis "
    "included 30 qualifying TMD-start units, 33 TMD-end units, and 32 "
    "TMD-center units. TMD-end-relative variance was directionally lower than "
    "the null median but did not reach permutation-based statistical "
    "significance (empirical one-sided P = 0.0829), whereas TMD-start and "
    "TMD-center analyses showed no evidence of reduced positional variance."
)

# prevent duplicate caption
already = any(
    p.text.strip().startswith(CAPTION_HEAD)
    for p in doc.paragraphs
)

if not already:

    target = None

    for p in doc.paragraphs:
        if p.text.strip() == "Data Availability":
            target = p
            break

    # fallback: insert before References
    if target is None:
        for p in doc.paragraphs:
            if p.text.strip() == "References":
                target = p
                break

    if target is None:
        raise RuntimeError(
            'Could not find "Data Availability" or "References".'
        )

    new_p = doc.add_paragraph(CAPTION)

    # preserve nearby body style
    try:
        new_p.style = target.style
    except Exception:
        pass

    target._element.addprevious(new_p._element)

    changes += 1


doc.save(DOCX)

print("Saved:", DOCX)
print("Changes:", changes)
