from docx import Document
from collections import Counter
from pathlib import Path

DOCX = Path(
    "/Users/liming/Desktop/秋山研/2026李明投稿论文/李明博论/"
    "MBE_submission_manuscript_v1.2.docx"
)

doc = Document(DOCX)

checks = {
    "Bacillales": [],
    "P = 0.0829": [],
    "739.08": [],
    "1,693.00": [],
    "30 TMD-start": [],
    "33 TMD-end": [],
    "32 TMD-center": [],
    "Supplementary Figure B1": [],
    "Supplementary Table B1": [],
    "Supplementary Table B2": [],
    "Supplementary Table B3": [],
    "Supplementary Table B4": [],
    "Supplementary Table B5": [],
    "Supplementary Table B6": [],
    "Supplementary Table B7": [],
    "Supplementary Table B8": [],
    "Supplementary Table B9": [],
}

for i, p in enumerate(doc.paragraphs):
    for key in checks:
        if key in p.text:
            checks[key].append(i)

print("=== KEY CONSISTENCY CHECK ===")
print()

for key, hits in checks.items():
    print(f"{key}: {len(hits)} hit(s) -> {hits}")


print()
print("=== POTENTIALLY OVERSTRONG WORDING ===")
print()

terms = [
    "replicated in Bacillales",
    "replication was successful",
    "universally conserved",
    "across bacteria",
    "general bacterial rule",
    "Bacillales confirmed",
]

found_any = False

for term in terms:
    for i, p in enumerate(doc.paragraphs):
        if term.lower() in p.text.lower():
            found_any = True
            print(f"[{i}] {term}")
            print(p.text)
            print()

if not found_any:
    print("No obvious overstrong wording found.")


print()
print("=== OUTDATED LIMITATION CHECK ===")
print()

old_phrases = [
    "the analysis was restricted to Enterobacterales",
    "the extent to which the same positional constraint applies across more deeply diverged bacterial lineages remains unknown",
]

for phrase in old_phrases:
    hits = [
        i for i, p in enumerate(doc.paragraphs)
        if phrase.lower() in p.text.lower()
    ]
    print(f"{phrase}: {len(hits)}")


print()
print("=== AVAILABILITY CHECK ===")
print()

for i, p in enumerate(doc.paragraphs):
    if (
        "Processed analysis tables underlying" in p.text
        or "Custom scripts used for ortholog filtering" in p.text
    ):
        print(f"[{i}]")
        print(p.text)
        print()


print()
print("=== FIGURE B1 CAPTION CHECK ===")
print()

for i, p in enumerate(doc.paragraphs):
    if p.text.startswith("Supplementary Figure B1."):
        print(f"[{i}]")
        print(repr(p.text))
        print()


print()
print("Audit complete.")
