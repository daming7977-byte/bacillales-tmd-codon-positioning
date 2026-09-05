from pathlib import Path
from docx import Document
import re

DOCX = Path(
    "/Users/liming/Desktop/秋山研/2026李明投稿论文/李明博论/"
    "MBE_submission_manuscript_v1.5_final_submission_candidate.docx"
)

doc = Document(DOCX)

changes = 0

# Fix known leftover
for p in doc.paragraphs:
    if "Supplementary Tables S14 and B9" in p.text:
        p.text = p.text.replace(
            "Supplementary Tables S14 and B9",
            "Supplementary Tables S14 and S15"
        )
        changes += 1

doc.save(DOCX)

# Reload for audit
doc = Document(DOCX)
text = "\n".join(p.text for p in doc.paragraphs)

print("Saved:", DOCX)
print("Corrections:", changes)

print()
print("=== ANY REMAINING B-NUMBER TOKENS ===")

hits = sorted(set(re.findall(r"\bB\d+\b", text)))

if hits:
    print(hits)
else:
    print("NONE")

print()
print("=== BACILLALES S REFERENCES ===")

for i in range(7, 16):
    n = len(re.findall(rf"\bS{i}\b", text))
    print(f"S{i}: {n}")

print()
print("Supplementary Figure S6:",
      text.count("Supplementary Figure S6"))

print()
print("=== TARGET PARAGRAPH ===")

for p in doc.paragraphs:
    if "full permutation-null summary" in p.text.lower():
        print(p.text)
        break
